#!/usr/bin/env python3
"""
Convert HTML file to Word document with proper formatting
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import re
import base64
from io import BytesIO

def clean_text(text):
    """Clean and normalize text"""
    if not text:
        return ""
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def add_heading(doc, text, level=1):
    """Add a heading to the document"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph to the document"""
    if not text or not text.strip():
        return None

    p = doc.add_paragraph()
    run = p.add_run(text)

    if bold:
        run.bold = True
    if italic:
        run.italic = True

    return p

def add_table_from_html(doc, html_table):
    """Convert HTML table to Word table"""
    rows = html_table.find_all('tr')
    if not rows:
        return

    # Count columns from first row
    first_row = rows[0]
    cols = len(first_row.find_all(['th', 'td']))

    if cols == 0:
        return

    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Light Grid Accent 1'

    for i, row in enumerate(rows):
        cells = row.find_all(['th', 'td'])
        for j, cell in enumerate(cells):
            if j < cols:
                table.rows[i].cells[j].text = clean_text(cell.get_text())

def convert_html_to_docx(html_path, output_path):
    """Convert HTML file to DOCX"""
    print(f"Reading HTML file: {html_path}")

    with open(html_path, 'r', encoding='utf-8') as f:
        html_content = f.read()

    print("Parsing HTML content...")
    soup = BeautifulSoup(html_content, 'lxml')

    # Create new Word document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Extract title
    title = soup.find('title')
    if title:
        doc.add_heading(clean_text(title.get_text()), 0)

    # Find main content (usually in body or main sections)
    body = soup.find('body')
    if not body:
        body = soup

    print("Converting content...")

    # Process elements
    for element in body.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div', 'table', 'ul', 'ol', 'pre', 'code']):

        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # Handle headings
            level = int(element.name[1])
            text = clean_text(element.get_text())
            if text:
                add_heading(doc, text, level=min(level, 9))

        elif element.name == 'p':
            # Handle paragraphs
            text = clean_text(element.get_text())
            if text:
                add_paragraph(doc, text)

        elif element.name == 'table':
            # Handle tables
            add_table_from_html(doc, element)
            doc.add_paragraph()  # Add space after table

        elif element.name in ['ul', 'ol']:
            # Handle lists
            for li in element.find_all('li', recursive=False):
                text = clean_text(li.get_text())
                if text:
                    p = doc.add_paragraph(text, style='List Bullet' if element.name == 'ul' else 'List Number')

        elif element.name in ['pre', 'code']:
            # Handle code blocks
            text = element.get_text()
            if text and text.strip():
                p = doc.add_paragraph(text)
                p.style = 'No Spacing'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)

    # Save document
    print(f"Saving document to: {output_path}")
    doc.save(output_path)
    print("Conversion complete!")

if __name__ == "__main__":
    html_file = "/Users/ugochi141/Desktop/M.S. Biotech/Fall 2025/Introduction to Bioinformatics/ACC/ACC.html"
    output_file = "/Users/ugochi141/Desktop/M.S. Biotech/Fall 2025/Introduction to Bioinformatics/ACC/ACC.docx"

    convert_html_to_docx(html_file, output_file)
